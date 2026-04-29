"""
Contract Red-lining Engine — AI-powered contract analysis with tracked changes in DOCX.
Reads a Word document, analyzes each clause, generates a modified DOCX with:
- Red-lined (strikethrough red) deletions
- Green inserted text for additions
- Margin comments explaining each change
"""
import io
import os
import re
import json
import logging
import aiohttp
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

logger = logging.getLogger("contract_redline")

GOOGLE_AI_KEY = os.environ.get("GOOGLE_AI_KEY", "")
ANTHROPIC_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
OPENAI_KEY = os.environ.get("OPENAI_KEY", "")

# OOXML namespace
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
nsmap = {'w': WORD_NS}


REDLINE_ANALYSIS_PROMPT = """You are an elite contract attorney at a top-tier Indian law firm. Analyze this contract clause-by-clause and identify EVERY issue that needs correction, improvement, or deletion.

For each issue, output a JSON object with:
- "clause_num": the clause/section number or identifier
- "original_text": the EXACT text from the contract that needs changing (must match verbatim)
- "action": "modify" | "delete" | "insert_after"
- "suggested_text": the replacement text (for modify), empty string (for delete), or new text to insert
- "reason": 1-2 sentence explanation of why this change is needed
- "severity": "critical" | "high" | "medium" | "low"
- "category": "risk" | "ambiguity" | "missing_protection" | "compliance" | "drafting_error" | "one_sided" | "market_standard"

ANALYSIS REQUIREMENTS:
1. Check for missing standard protections (indemnity caps, limitation of liability, force majeure, dispute resolution)
2. Flag one-sided clauses that favor only one party
3. Identify ambiguous language that could lead to disputes
4. Check limitation periods and notice periods against Indian law standards
5. Verify compliance with Indian Contract Act, 1872; Specific Relief Act, 1963; Arbitration Act, 1996
6. Flag missing governing law/jurisdiction clauses
7. Check for proper defined terms usage
8. Identify overly broad non-compete/non-solicitation clauses
9. Flag unlimited liability provisions
10. Check termination clauses for fairness

OUTPUT FORMAT — Return ONLY a JSON array of suggestion objects. No other text.
Example:
[
  {
    "clause_num": "3.2",
    "original_text": "The Contractor shall be liable for all damages without any limitation.",
    "action": "modify",
    "suggested_text": "The Contractor's aggregate liability under this Agreement shall not exceed the total fees paid by the Client in the twelve (12) months preceding the claim.",
    "reason": "Unlimited liability is commercially unreasonable and could expose the Contractor to disproportionate risk. A liability cap aligned with contract value is market standard.",
    "severity": "critical",
    "category": "risk"
  }
]

CONTRACT TEXT:
"""


def extract_docx_paragraphs(docx_bytes: bytes) -> list:
    """Extract paragraphs from DOCX with their text and style info."""
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "index": i,
                "text": text,
                "style": para.style.name if para.style else "Normal",
                "alignment": str(para.alignment) if para.alignment else "LEFT",
            })
    return paragraphs


def extract_full_text(docx_bytes: bytes) -> str:
    """Extract full text from DOCX."""
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


async def analyze_contract_for_redlines(contract_text: str, focus_area: str = "") -> list:
    """Send contract text to AI for clause-by-clause analysis. Returns list of suggestions."""
    prompt = REDLINE_ANALYSIS_PROMPT + contract_text[:80000]
    if focus_area:
        prompt += f"\n\nSPECIAL FOCUS: Pay extra attention to {focus_area}"

    suggestions = []

    async with aiohttp.ClientSession() as session:
        # Try Gemini first
        if GOOGLE_AI_KEY:
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={GOOGLE_AI_KEY}"
                payload = {
                    "contents": [{"role": "user", "parts": [{"text": prompt}]}],
                    "generationConfig": {
                        "temperature": 0.1,
                        "maxOutputTokens": 16384,
                        "responseMimeType": "application/json"
                    }
                }
                async with session.post(url, json=payload, timeout=aiohttp.ClientTimeout(total=120)) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        parts = data.get("candidates", [{}])[0].get("content", {}).get("parts", [])
                        text = "".join([p.get("text", "") for p in parts if "text" in p])
                        suggestions = _parse_suggestions(text)
                        if suggestions:
                            return suggestions
            except Exception as e:
                logger.warning(f"Gemini redline analysis failed: {e}")

        # Fallback to Claude
        if ANTHROPIC_KEY:
            try:
                payload = {
                    "model": "claude-sonnet-4-20250514",
                    "max_tokens": 16384,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.1,
                }
                async with session.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": ANTHROPIC_KEY,
                        "anthropic-version": "2023-06-01",
                        "Content-Type": "application/json"
                    },
                    json=payload, timeout=aiohttp.ClientTimeout(total=120)
                ) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        text = "\n".join([b["text"] for b in data.get("content", []) if b.get("type") == "text"])
                        suggestions = _parse_suggestions(text)
                        if suggestions:
                            return suggestions
            except Exception as e:
                logger.warning(f"Claude redline analysis failed: {e}")

        # Fallback to GPT-4o
        if OPENAI_KEY:
            try:
                payload = {
                    "model": "gpt-4o",
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.1,
                    "max_tokens": 16384,
                    "response_format": {"type": "json_object"}
                }
                async with session.post(
                    "https://api.openai.com/v1/chat/completions",
                    headers={"Authorization": f"Bearer {OPENAI_KEY}", "Content-Type": "application/json"},
                    json=payload, timeout=aiohttp.ClientTimeout(total=90)
                ) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        text = data["choices"][0]["message"]["content"]
                        suggestions = _parse_suggestions(text)
                        if suggestions:
                            return suggestions
            except Exception as e:
                logger.warning(f"GPT-4o redline analysis failed: {e}")

    return suggestions


def _parse_suggestions(text: str) -> list:
    """Parse AI response into structured suggestions list."""
    text = text.strip()
    # Try to extract JSON array from the response
    try:
        # Direct parse
        result = json.loads(text)
        if isinstance(result, list):
            return _validate_suggestions(result)
        if isinstance(result, dict):
            # GPT-4o sometimes wraps in {"suggestions": [...]}
            for key in ["suggestions", "changes", "redlines", "issues", "results"]:
                if key in result and isinstance(result[key], list):
                    return _validate_suggestions(result[key])
    except json.JSONDecodeError:
        pass

    # Try to find JSON array in the text
    match = re.search(r'\[[\s\S]*\]', text)
    if match:
        try:
            result = json.loads(match.group())
            if isinstance(result, list):
                return _validate_suggestions(result)
        except json.JSONDecodeError:
            pass

    logger.warning("Could not parse AI suggestions as JSON")
    return []


def _validate_suggestions(suggestions: list) -> list:
    """Validate and clean suggestion objects."""
    valid = []
    for s in suggestions:
        if not isinstance(s, dict):
            continue
        if not s.get("original_text") and s.get("action") != "insert_after":
            continue
        valid.append({
            "clause_num": s.get("clause_num", ""),
            "original_text": s.get("original_text", ""),
            "action": s.get("action", "modify"),
            "suggested_text": s.get("suggested_text", ""),
            "reason": s.get("reason", ""),
            "severity": s.get("severity", "medium"),
            "category": s.get("category", "risk"),
        })
    return valid


def generate_redlined_docx(original_bytes: bytes, suggestions: list) -> bytes:
    """Generate a new DOCX with visual red-lining (strikethrough + colored text).

    Uses visual formatting since python-docx doesn't support OOXML tracked changes natively:
    - Deleted text: Red strikethrough
    - Added text: Blue/green underlined
    - Each change preceded by a comment-style annotation
    """
    doc = Document(io.BytesIO(original_bytes))

    # Build a lookup of original_text -> suggestion for quick matching
    suggestion_map = {}
    for s in suggestions:
        if s.get("original_text"):
            suggestion_map[s["original_text"].strip()] = s

    # Track which suggestions were applied
    applied = set()

    # Process each paragraph
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue

        for orig_text, suggestion in suggestion_map.items():
            if orig_text in applied:
                continue

            if orig_text in para_text:
                applied.add(orig_text)
                action = suggestion.get("action", "modify")
                suggested = suggestion.get("suggested_text", "")
                reason = suggestion.get("reason", "")
                severity = suggestion.get("severity", "medium")
                clause = suggestion.get("clause_num", "")

                # Clear existing runs
                for run in para.runs:
                    run.text = ""

                # Rebuild paragraph with red-lined changes
                if action == "delete":
                    # Show deleted text in red with strikethrough
                    _add_annotation_run(para, f"[DELETED — {reason}]", severity)
                    deleted_run = para.add_run(para_text)
                    deleted_run.font.color.rgb = RGBColor(220, 38, 38)  # Red
                    deleted_run.font.strike = True
                    deleted_run.font.size = Pt(10)

                elif action == "modify":
                    # Replace the matched portion
                    before = para_text[:para_text.index(orig_text)]
                    after = para_text[para_text.index(orig_text) + len(orig_text):]

                    # Text before the change (unchanged)
                    if before:
                        unchanged = para.add_run(before)
                        unchanged.font.size = Pt(11)

                    # Show deleted text (red strikethrough)
                    deleted_run = para.add_run(orig_text)
                    deleted_run.font.color.rgb = RGBColor(220, 38, 38)
                    deleted_run.font.strike = True
                    deleted_run.font.size = Pt(11)

                    # Show added text (blue)
                    if suggested:
                        added_run = para.add_run(" " + suggested)
                        added_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue
                        added_run.font.underline = True
                        added_run.font.size = Pt(11)

                    # Text after the change (unchanged)
                    if after:
                        unchanged = para.add_run(after)
                        unchanged.font.size = Pt(11)

                    # Add reason as annotation below
                    _add_annotation_run(para, f" [{severity.upper()}: {reason}]", severity)

                elif action == "insert_after":
                    # Keep original text
                    unchanged = para.add_run(para_text)
                    unchanged.font.size = Pt(11)

                break  # Only apply first matching suggestion per paragraph

    # Add a summary page at the end
    _add_summary_page(doc, suggestions, applied)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _add_annotation_run(para, text, severity):
    """Add a small colored annotation run to a paragraph."""
    colors = {
        "critical": RGBColor(185, 28, 28),
        "high": RGBColor(217, 119, 6),
        "medium": RGBColor(37, 99, 235),
        "low": RGBColor(75, 85, 99),
    }
    run = para.add_run(text)
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = colors.get(severity, RGBColor(107, 114, 128))


def _add_summary_page(doc, suggestions, applied):
    """Add a change summary page at the end of the document."""
    doc.add_page_break()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CONTRACT REVIEW SUMMARY")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(10, 10, 10)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Generated by Spectr AI — {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph("")  # spacer

    # Stats
    stats_para = doc.add_paragraph()
    total = len(suggestions)
    critical = sum(1 for s in suggestions if s.get("severity") == "critical")
    high = sum(1 for s in suggestions if s.get("severity") == "high")
    medium = sum(1 for s in suggestions if s.get("severity") == "medium")
    low = sum(1 for s in suggestions if s.get("severity") == "low")

    stats_run = stats_para.add_run(f"Total Issues Found: {total}  |  Critical: {critical}  |  High: {high}  |  Medium: {medium}  |  Low: {low}")
    stats_run.font.size = Pt(11)
    stats_run.font.bold = True

    doc.add_paragraph("")  # spacer

    # Detailed list
    for i, s in enumerate(suggestions, 1):
        severity_label = s.get("severity", "medium").upper()
        category_label = s.get("category", "").replace("_", " ").title()
        clause = s.get("clause_num", "N/A")

        # Issue header
        header = doc.add_paragraph()
        header_run = header.add_run(f"{i}. [{severity_label}] Clause {clause} — {category_label}")
        header_run.font.size = Pt(11)
        header_run.font.bold = True

        severity_colors = {
            "CRITICAL": RGBColor(185, 28, 28),
            "HIGH": RGBColor(217, 119, 6),
            "MEDIUM": RGBColor(37, 99, 235),
            "LOW": RGBColor(75, 85, 99),
        }
        header_run.font.color.rgb = severity_colors.get(severity_label, RGBColor(0, 0, 0))

        # Original text
        if s.get("original_text"):
            orig = doc.add_paragraph()
            orig_label = orig.add_run("Original: ")
            orig_label.font.size = Pt(10)
            orig_label.font.bold = True
            orig_text = orig.add_run(s["original_text"][:300])
            orig_text.font.size = Pt(10)
            orig_text.font.color.rgb = RGBColor(220, 38, 38)
            orig_text.font.strike = True

        # Suggested text
        action = s.get("action", "modify")
        if action == "modify" and s.get("suggested_text"):
            sugg = doc.add_paragraph()
            sugg_label = sugg.add_run("Suggested: ")
            sugg_label.font.size = Pt(10)
            sugg_label.font.bold = True
            sugg_text = sugg.add_run(s["suggested_text"][:300])
            sugg_text.font.size = Pt(10)
            sugg_text.font.color.rgb = RGBColor(37, 99, 235)
            sugg_text.font.underline = True
        elif action == "delete":
            sugg = doc.add_paragraph()
            sugg_run = sugg.add_run("Recommendation: DELETE this clause entirely")
            sugg_run.font.size = Pt(10)
            sugg_run.font.bold = True
            sugg_run.font.color.rgb = RGBColor(220, 38, 38)

        # Reason
        if s.get("reason"):
            reason = doc.add_paragraph()
            reason_label = reason.add_run("Reason: ")
            reason_label.font.size = Pt(10)
            reason_label.font.bold = True
            reason_text = reason.add_run(s["reason"])
            reason_text.font.size = Pt(10)
            reason_text.font.italic = True
            reason_text.font.color.rgb = RGBColor(75, 85, 99)

        doc.add_paragraph("")  # spacer between issues


async def process_contract_redline(docx_bytes: bytes, focus_area: str = "") -> dict:
    """Full pipeline: extract text → AI analysis → generate red-lined DOCX.

    Returns dict with:
    - redlined_docx: bytes of the modified DOCX
    - suggestions: list of suggestion dicts
    - stats: summary statistics
    """
    # Step 1: Extract text
    full_text = extract_full_text(docx_bytes)
    if not full_text or len(full_text) < 50:
        return {"error": "Could not extract meaningful text from the document."}

    logger.info(f"Contract text extracted: {len(full_text)} chars")

    # Step 2: AI analysis
    suggestions = await analyze_contract_for_redlines(full_text, focus_area)
    if not suggestions:
        return {
            "error": "AI analysis did not produce any suggestions. The contract may already be well-drafted.",
            "suggestions": [],
            "stats": {"total": 0}
        }

    logger.info(f"AI produced {len(suggestions)} suggestions")

    # Step 3: Generate red-lined DOCX
    redlined_bytes = generate_redlined_docx(docx_bytes, suggestions)

    # Step 4: Compile stats
    stats = {
        "total": len(suggestions),
        "critical": sum(1 for s in suggestions if s.get("severity") == "critical"),
        "high": sum(1 for s in suggestions if s.get("severity") == "high"),
        "medium": sum(1 for s in suggestions if s.get("severity") == "medium"),
        "low": sum(1 for s in suggestions if s.get("severity") == "low"),
        "by_category": {},
    }
    for s in suggestions:
        cat = s.get("category", "other")
        stats["by_category"][cat] = stats["by_category"].get(cat, 0) + 1

    return {
        "redlined_docx": redlined_bytes,
        "suggestions": suggestions,
        "stats": stats,
    }

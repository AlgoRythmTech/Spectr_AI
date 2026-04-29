"""
Enhanced Document Export
Professional legal document formatting with proper cause titles,
paragraph numbering, court header styles, signature blocks, and stamp paper templates.
"""
import os
import io
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

# Court-specific formatting templates
COURT_FORMATS = {
    "supreme_court": {
        "header": "IN THE SUPREME COURT OF INDIA",
        "sub_header": "CIVIL/CRIMINAL APPELLATE/ORIGINAL JURISDICTION",
        "font": "Times New Roman",
        "font_size": 13,
        "line_spacing": 1.5,
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_left": 3.8,
        "margin_right": 2.5,
    },
    "high_court": {
        "header": "IN THE HIGH COURT OF {state}",
        "sub_header": "{bench} BENCH",
        "font": "Times New Roman",
        "font_size": 13,
        "line_spacing": 1.5,
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_left": 3.8,
        "margin_right": 2.5,
    },
    "district_court": {
        "header": "IN THE COURT OF {judge_designation}",
        "sub_header": "{district}, {state}",
        "font": "Times New Roman",
        "font_size": 13,
        "line_spacing": 1.5,
        "margin_top": 2.5,
        "margin_bottom": 2.0,
        "margin_left": 3.5,
        "margin_right": 2.5,
    },
    "nclt": {
        "header": "NATIONAL COMPANY LAW TRIBUNAL",
        "sub_header": "{bench} BENCH",
        "font": "Times New Roman",
        "font_size": 13,
        "line_spacing": 1.5,
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_left": 3.5,
        "margin_right": 2.5,
    },
    "gst_tribunal": {
        "header": "GOODS AND SERVICES TAX APPELLATE TRIBUNAL",
        "sub_header": "{bench} BENCH",
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.5,
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_left": 3.5,
        "margin_right": 2.5,
    },
    "tax_authority": {
        "header": "BEFORE THE {authority}",
        "sub_header": "",
        "font": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.5,
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_left": 2.5,
        "margin_right": 2.5,
    }
}

def generate_court_document(
    content: str,
    court_type: str = "high_court",
    case_number: str = "",
    petitioner: str = "",
    respondent: str = "",
    document_type: str = "PETITION",
    advocate_name: str = "",
    advocate_enrollment: str = "",
    court_params: dict = None,
    sections: list = None
) -> bytes:
    """Generate a professionally formatted court document as DOCX bytes."""
    
    doc = Document()
    fmt = COURT_FORMATS.get(court_type, COURT_FORMATS["high_court"])
    params = court_params or {}
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(fmt["margin_top"])
        section.bottom_margin = Cm(fmt["margin_bottom"])
        section.left_margin = Cm(fmt["margin_left"])
        section.right_margin = Cm(fmt["margin_right"])
    
    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = fmt["font"]
    font.size = Pt(fmt["font_size"])
    
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = fmt["line_spacing"]
    
    # === COURT HEADER ===
    header_text = fmt["header"]
    for key, val in params.items():
        header_text = header_text.replace(f"{{{key}}}", str(val))
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header_text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = fmt["font"]
    
    # Sub header
    if fmt["sub_header"]:
        sub_text = fmt["sub_header"]
        for key, val in params.items():
            sub_text = sub_text.replace(f"{{{key}}}", str(val))
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(sub_text)
        run2.font.size = Pt(12)
        run2.font.name = fmt["font"]
    
    # Case number
    if case_number:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(f"\n{case_number}")
        run3.bold = True
        run3.font.size = Pt(12)
    
    # === CAUSE TITLE ===
    doc.add_paragraph()  # Spacer
    
    if petitioner:
        # Petitioner
        p_pet = doc.add_paragraph()
        p_pet.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_pet = p_pet.add_run(petitioner.upper())
        run_pet.bold = True
        run_pet.font.size = Pt(12)
        
        # ... Petitioner/Appellant label
        p_label1 = doc.add_paragraph()
        p_label1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_l1 = p_label1.add_run("...Petitioner/Appellant")
        run_l1.italic = True
        
        # Versus
        p_vs = doc.add_paragraph()
        p_vs.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_vs = p_vs.add_run("VERSUS")
        run_vs.bold = True
        run_vs.font.size = Pt(12)
        
        # Respondent
        p_resp = doc.add_paragraph()
        p_resp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_resp = p_resp.add_run(respondent.upper())
        run_resp.bold = True
        run_resp.font.size = Pt(12)
        
        # ... Respondent label
        p_label2 = doc.add_paragraph()
        p_label2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_l2 = p_label2.add_run("...Respondent")
        run_l2.italic = True
    
    # Divider
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_div = p_div.add_run("_" * 50)
    run_div.font.color.rgb = RGBColor(150, 150, 150)
    
    # Document type heading
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_type = p_type.add_run(f"\n{document_type.upper()}")
    run_type.bold = True
    run_type.font.size = Pt(14)
    run_type.underline = True
    
    doc.add_paragraph()  # Spacer
    
    # === BODY CONTENT ===
    if sections:
        for sec in sections:
            if sec.get("title"):
                p_title = doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run_title = p_title.add_run(sec["title"].upper())
                run_title.bold = True
                run_title.underline = True
                run_title.font.size = Pt(12)
            
            if sec.get("content"):
                _add_formatted_content(doc, sec["content"], fmt)
    elif content:
        _add_formatted_content(doc, content, fmt)
    
    # === PRAYER CLAUSE (if present in content) ===
    # (Already embedded in content)
    
    # === SIGNATURE BLOCK ===
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_place = doc.add_paragraph()
    p_place.add_run(f"Place: ________________")
    p_date = doc.add_paragraph()
    p_date.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
    
    doc.add_paragraph()
    
    if advocate_name:
        p_adv = doc.add_paragraph()
        p_adv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_adv = p_adv.add_run(f"Respectfully submitted,\n\n\n{advocate_name}")
        run_adv.font.size = Pt(12)
        
        if advocate_enrollment:
            p_enr = doc.add_paragraph()
            p_enr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_enr = p_enr.add_run(f"Enrollment No.: {advocate_enrollment}")
            run_enr.font.size = Pt(10)
    else:
        p_adv = doc.add_paragraph()
        p_adv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_adv = p_adv.add_run("Respectfully submitted,\n\n\nAdvocate for the Petitioner/Appellant")
    
    # === VERIFICATION (for affidavits/petitions) ===
    doc.add_page_break()
    p_ver_title = doc.add_paragraph()
    p_ver_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ver = p_ver_title.add_run("VERIFICATION")
    run_ver.bold = True
    run_ver.underline = True
    run_ver.font.size = Pt(13)
    
    p_ver = doc.add_paragraph()
    p_ver.add_run(
        f"I, {petitioner or '________________'}, the petitioner/appellant above named, "
        f"do hereby verify that the contents of the above petition/appeal are true and correct "
        f"to my knowledge and belief and nothing material has been concealed therefrom."
    )
    
    p_ver2 = doc.add_paragraph()
    p_ver2.add_run(f"\nVerified at ______________ on this {datetime.now().strftime('%d')} day of {datetime.now().strftime('%B, %Y')}.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.add_run("DEPONENT")
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _add_formatted_content(doc, content: str, fmt: dict):
    """Add content to document with proper legal paragraph numbering."""
    lines = content.split('\n')
    para_num = 0
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        
        # Check if it's already numbered
        import re
        num_match = re.match(r'^(\d+)\.\s*', stripped)
        
        if num_match:
            # Use existing numbering
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.27)
            run = p.add_run(stripped)
            run.font.name = fmt["font"]
            run.font.size = Pt(fmt["font_size"])
        elif stripped.startswith('- ') or stripped.startswith('* '):
            # Bullet point
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            run = p.add_run("• " + stripped[2:])
            run.font.name = fmt["font"]
            run.font.size = Pt(fmt["font_size"])
        elif stripped.startswith('**') and stripped.endswith('**'):
            # Bold heading
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip('*'))
            run.bold = True
            run.font.name = fmt["font"]
            run.font.size = Pt(fmt["font_size"])
        else:
            # Regular paragraph with auto-numbering for substantial paragraphs
            if len(stripped) > 50:
                para_num += 1
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(1.27)
                run = p.add_run(f"{para_num}.\t{stripped}")
            else:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
            
            run.font.name = fmt["font"]
            run.font.size = Pt(fmt["font_size"])


def generate_tax_notice_reply(
    notice_details: dict,
    reply_content: str,
    authority_type: str = "gst"
) -> bytes:
    """Generate a formatted reply to a tax notice."""
    
    if authority_type == "gst":
        court_type = "tax_authority"
        params = {"authority": "THE ASSISTANT/DEPUTY/JOINT COMMISSIONER OF CGST"}
        doc_type = "REPLY TO SHOW CAUSE NOTICE"
    else:
        court_type = "tax_authority"
        params = {"authority": "THE ASSESSING OFFICER, INCOME TAX DEPARTMENT"}
        doc_type = "REPLY TO NOTICE UNDER SECTION " + notice_details.get("section", "148")
    
    return generate_court_document(
        content=reply_content,
        court_type=court_type,
        case_number=notice_details.get("notice_number", ""),
        petitioner=notice_details.get("taxpayer_name", ""),
        respondent=notice_details.get("authority_name", ""),
        document_type=doc_type,
        advocate_name=notice_details.get("advocate_name", ""),
        advocate_enrollment=notice_details.get("enrollment", ""),
        court_params=params
    )


def generate_memo_export(content: str, title: str = "Legal Memorandum") -> bytes:
    """Generate a professional legal memo."""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = "Georgia"
    style.font.size = Pt(11.5)
    style.paragraph_format.line_spacing = 1.5
    
    # Header
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = p_header.add_run("CONFIDENTIAL — PRIVILEGED AND CONFIDENTIAL")
    run_h.font.size = Pt(9)
    run_h.font.color.rgb = RGBColor(150, 0, 0)
    run_h.bold = True
    
    # Line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.add_run("_" * 60).font.color.rgb = RGBColor(200, 200, 200)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run(f"\n{title.upper()}")
    run_t.bold = True
    run_t.font.size = Pt(16)
    
    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = p_date.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Body
    _add_formatted_content(doc, content, {"font": "Georgia", "font_size": 11.5})
    
    # Footer
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run("\n— Generated by Spectr Intelligence Platform —")
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(150, 150, 150)
    run_f.italic = True
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

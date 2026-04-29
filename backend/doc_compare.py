"""
Document Comparison Engine — Spectr Grade
Diffs two document versions, identifies changes, and produces a comparison report.
Supports: DOCX, PDF, TXT. Outputs: JSON diff, comparison DOCX with tracked changes.
"""
import io
import os
import re
import json
import logging
import difflib
import aiohttp
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("doc_compare")

GOOGLE_AI_KEY = os.environ.get("GOOGLE_AI_KEY", "")
ANTHROPIC_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
GROQ_KEY = os.environ.get("GROQ_KEY", "")


def extract_text_from_bytes(file_data: bytes, filename: str) -> str:
    """Extract text from DOCX, PDF, or TXT."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext in ("docx", "doc"):
        doc = Document(io.BytesIO(file_data))
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext == "pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_data))
            return "\n".join([page.extract_text() or "" for page in reader.pages])
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""
    elif ext in ("txt", "text", "md"):
        return file_data.decode("utf-8", errors="ignore")
    return ""


def compute_diff(text_a: str, text_b: str) -> dict:
    """Compute a structured diff between two texts.

    Returns:
        dict with: changes (list), summary stats, similarity_pct
    """
    lines_a = text_a.splitlines()
    lines_b = text_b.splitlines()

    differ = difflib.unified_diff(lines_a, lines_b, lineterm="", n=2)
    diff_lines = list(differ)

    # Parse unified diff into structured changes
    changes = []
    current_change = None

    for line in diff_lines:
        if line.startswith("@@"):
            if current_change:
                changes.append(current_change)
            # Parse line numbers
            match = re.search(r'@@ -(\d+),?\d* \+(\d+),?\d* @@', line)
            old_line = int(match.group(1)) if match else 0
            new_line = int(match.group(2)) if match else 0
            current_change = {
                "old_line": old_line,
                "new_line": new_line,
                "removed": [],
                "added": [],
                "context": [],
            }
        elif line.startswith("---") or line.startswith("+++"):
            continue
        elif current_change is not None:
            if line.startswith("-"):
                current_change["removed"].append(line[1:])
            elif line.startswith("+"):
                current_change["added"].append(line[1:])
            else:
                current_change["context"].append(line[1:] if line.startswith(" ") else line)

    if current_change:
        changes.append(current_change)

    # Compute similarity
    seq = difflib.SequenceMatcher(None, text_a, text_b)
    similarity = round(seq.ratio() * 100, 1)

    # Classify changes
    total_added = sum(len(c["added"]) for c in changes)
    total_removed = sum(len(c["removed"]) for c in changes)

    return {
        "changes": changes,
        "total_changes": len(changes),
        "lines_added": total_added,
        "lines_removed": total_removed,
        "similarity_pct": similarity,
        "doc_a_lines": len(lines_a),
        "doc_b_lines": len(lines_b),
    }


async def smart_compare(text_a: str, text_b: str) -> dict:
    """AI-powered comparison that explains the significance of changes.
    Goes beyond raw diff — identifies legal/business impact.
    """
    diff_result = compute_diff(text_a, text_b)

    # Build a summary of changes for AI analysis
    change_summary = []
    for i, c in enumerate(diff_result["changes"][:30]):  # Limit to 30 changes
        removed = " ".join(c["removed"])[:200]
        added = " ".join(c["added"])[:200]
        if removed or added:
            change_summary.append(f"Change {i+1}:\n  Removed: {removed}\n  Added: {added}")

    if not change_summary:
        diff_result["ai_analysis"] = "No significant changes detected between the two versions."
        return diff_result

    prompt = f"""Analyze these changes between two versions of a legal/business document.
For each significant change, explain:
1. What was changed
2. The legal or business significance
3. Whether the change favours Party A or Party B (or is neutral)
4. Risk level: LOW, MEDIUM, HIGH, CRITICAL

Focus on: liability changes, obligation changes, term changes, financial impact, compliance impact.

CHANGES:
{chr(10).join(change_summary[:20])}

Return a JSON array of analysis objects:
[{{"change_num": 1, "description": "...", "significance": "...", "favours": "Party A|Party B|Neutral", "risk": "LOW|MEDIUM|HIGH|CRITICAL"}}]
"""

    # Quick AI analysis
    analysis = await _quick_ai_call(prompt)
    diff_result["ai_analysis"] = analysis

    return diff_result


def generate_comparison_docx(text_a: str, text_b: str, name_a: str = "Version A", name_b: str = "Version B") -> bytes:
    """Generate a Word document showing the comparison with color-coded changes."""
    doc = Document()

    # Styling
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOCUMENT COMPARISON REPORT")
    run.bold = True
    run.font.size = Pt(16)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"Generated by Spectr AI — {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph()

    # Legend
    legend = doc.add_paragraph()
    del_run = legend.add_run("Red strikethrough = Removed from " + name_a + "    ")
    del_run.font.color.rgb = RGBColor(220, 38, 38)
    del_run.font.strike = True
    del_run.font.size = Pt(10)

    add_run = legend.add_run("Blue underline = Added in " + name_b)
    add_run.font.color.rgb = RGBColor(37, 99, 235)
    add_run.font.underline = True
    add_run.font.size = Pt(10)

    doc.add_paragraph()

    # Compute diff
    lines_a = text_a.splitlines()
    lines_b = text_b.splitlines()

    opcodes = difflib.SequenceMatcher(None, lines_a, lines_b).get_opcodes()

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            for line in lines_a[i1:i2]:
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line)
                    r.font.size = Pt(11)
        elif tag == "replace":
            # Show removed (red strikethrough) then added (blue underline)
            for line in lines_a[i1:i2]:
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line)
                    r.font.color.rgb = RGBColor(220, 38, 38)
                    r.font.strike = True
                    r.font.size = Pt(11)
            for line in lines_b[j1:j2]:
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line)
                    r.font.color.rgb = RGBColor(37, 99, 235)
                    r.font.underline = True
                    r.font.size = Pt(11)
        elif tag == "delete":
            for line in lines_a[i1:i2]:
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line)
                    r.font.color.rgb = RGBColor(220, 38, 38)
                    r.font.strike = True
                    r.font.size = Pt(11)
        elif tag == "insert":
            for line in lines_b[j1:j2]:
                if line.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(line)
                    r.font.color.rgb = RGBColor(37, 99, 235)
                    r.font.underline = True
                    r.font.size = Pt(11)

    # Summary stats
    doc.add_page_break()
    summary_title = doc.add_paragraph()
    summary_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = summary_title.add_run("COMPARISON SUMMARY")
    sr.bold = True
    sr.font.size = Pt(14)

    diff = compute_diff(text_a, text_b)
    stats = doc.add_paragraph()
    stats.add_run(f"Similarity: {diff['similarity_pct']}%\n").bold = True
    stats.add_run(f"Lines added: {diff['lines_added']}\n")
    stats.add_run(f"Lines removed: {diff['lines_removed']}\n")
    stats.add_run(f"Change blocks: {diff['total_changes']}\n")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def _quick_ai_call(prompt: str) -> str:
    """Quick AI call for change analysis."""
    # Try Gemini Flash first (fast)
    if GOOGLE_AI_KEY:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={GOOGLE_AI_KEY}"
            payload = {
                "contents": [{"role": "user", "parts": [{"text": prompt}]}],
                "generationConfig": {"temperature": 0.1, "maxOutputTokens": 4096}
            }
            async with aiohttp.ClientSession() as session:
                async with session.post(url, json=payload,
                                        timeout=aiohttp.ClientTimeout(total=30)) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        parts = data.get("candidates", [{}])[0].get("content", {}).get("parts", [])
                        return "".join([p.get("text", "") for p in parts if "text" in p])
        except Exception as e:
            logger.warning(f"Gemini comparison analysis failed: {e}")

    # Fallback to Groq
    if GROQ_KEY:
        try:
            async with aiohttp.ClientSession() as session:
                payload = {
                    "model": "llama-3.3-70b-versatile",
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.1,
                    "max_tokens": 4096
                }
                async with session.post(
                    "https://api.groq.com/openai/v1/chat/completions",
                    headers={"Authorization": f"Bearer {GROQ_KEY}", "Content-Type": "application/json"},
                    json=payload, timeout=aiohttp.ClientTimeout(total=20)
                ) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        return data["choices"][0]["message"]["content"]
        except Exception as e:
            logger.warning(f"Groq comparison analysis failed: {e}")

    return "AI analysis unavailable."

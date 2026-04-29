"""
Smart Sheet Sampler — generates a condensed "data profile" for any size sheet.

Problem: a 50-col × 10K-row sheet = ~500K cells = millions of tokens if dumped.
Solution: extract a compact profile (~300-800 tokens) with everything the LLM needs
to reason about the data without seeing every cell:
  - Per-column: name, detected type, unique count, nulls, sample values, range/stats
  - Headers + first N rows + last N rows
  - Detected structure: merged cells, formulas, named ranges, multiple header rows
  - Cross-sheet relationships: which sheets reference which

Used to ground the LLM planner so it writes correct Python on first attempt
without needing to load all data into context.
"""
import io
import re
import json
import logging
from collections import Counter
from typing import Optional

logger = logging.getLogger("sheet_profiler")


# === TYPE DETECTION ===

_DATE_PATTERNS = [
    r'^\d{4}[-/]\d{1,2}[-/]\d{1,2}$',             # 2026-01-15 or 2026/01/15
    r'^\d{1,2}[-/]\d{1,2}[-/]\d{4}$',             # 15-01-2026
    r'^\d{1,2}[-/\s](?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[-/\s]\d{2,4}$',
]
_CURRENCY_PATTERNS = [
    r'^(?:₹|Rs\.?|INR)\s*[\d,]+(?:\.\d+)?\s*(?:crore|lakh|cr|L)?$',
    r'^\$\s*[\d,]+(?:\.\d+)?$',
]
_GSTIN_PATTERN = re.compile(r'^\d{2}[A-Z]{5}\d{4}[A-Z]\d[A-Z0-9]{2}$')
_PAN_PATTERN = re.compile(r'^[A-Z]{5}\d{4}[A-Z]$')
_EMAIL_PATTERN = re.compile(r'^[^\s@]+@[^\s@]+\.[^\s@]+$')
_PHONE_PATTERN = re.compile(r'^\+?\d[\d\s\-\(\)]{8,14}$')


def _detect_cell_type(val) -> str:
    """Classify a cell value into a semantic type."""
    if val is None or val == "":
        return "empty"
    if isinstance(val, bool):
        return "boolean"
    if isinstance(val, (int, float)):
        return "number"
    # Check if it's a datetime object
    if hasattr(val, 'year') and hasattr(val, 'month'):
        return "date"
    s = str(val).strip()
    if not s:
        return "empty"
    if s.startswith("="):
        return "formula"
    # Try numeric (after stripping commas, currency)
    cleaned = re.sub(r'[,₹$\s]|(Rs\.?|INR|crore|lakh|cr|L)', '', s, flags=re.IGNORECASE)
    try:
        float(cleaned)
        for pat in _CURRENCY_PATTERNS:
            if re.match(pat, s, re.IGNORECASE):
                return "currency"
        return "number_str"
    except ValueError:
        pass
    # Date patterns
    for pat in _DATE_PATTERNS:
        if re.match(pat, s, re.IGNORECASE):
            return "date_str"
    # Indian identifiers
    if _GSTIN_PATTERN.match(s):
        return "gstin"
    if _PAN_PATTERN.match(s):
        return "pan"
    if _EMAIL_PATTERN.match(s):
        return "email"
    if _PHONE_PATTERN.match(s):
        return "phone"
    # Percentage
    if re.match(r'^\d+(?:\.\d+)?\s*%$', s):
        return "percentage"
    return "text"


def _column_profile(values: list, col_idx: int, header: str = None) -> dict:
    """Analyze a single column. values is a list of cell values (excluding header)."""
    non_empty = [v for v in values if v not in (None, "")]
    total = len(values)
    nulls = total - len(non_empty)
    unique_vals = set()
    for v in non_empty:
        try:
            unique_vals.add(str(v)[:60])
        except Exception:
            pass
    unique_count = len(unique_vals)

    # Type distribution
    type_counts = Counter()
    numeric_vals = []
    for v in non_empty:
        t = _detect_cell_type(v)
        type_counts[t] += 1
        if t in ("number", "number_str", "currency"):
            try:
                cleaned = re.sub(r'[,₹$\s]|(Rs\.?|INR|crore|lakh|cr|L)', '', str(v), flags=re.IGNORECASE)
                numeric_vals.append(float(cleaned))
            except Exception:
                pass

    dominant_type = type_counts.most_common(1)[0][0] if type_counts else "empty"

    # Numeric stats
    stats = {}
    if numeric_vals:
        numeric_vals.sort()
        stats["min"] = numeric_vals[0]
        stats["max"] = numeric_vals[-1]
        stats["sum"] = sum(numeric_vals)
        stats["mean"] = sum(numeric_vals) / len(numeric_vals) if numeric_vals else 0
        # Median
        n = len(numeric_vals)
        stats["median"] = numeric_vals[n // 2] if n % 2 else (numeric_vals[n//2 - 1] + numeric_vals[n//2]) / 2

    # Sample values (first 3 unique non-empty)
    samples = []
    seen = set()
    for v in non_empty:
        s = str(v)[:60]
        if s not in seen:
            seen.add(s)
            samples.append(s)
            if len(samples) >= 3:
                break

    return {
        "column": col_idx,
        "header": header or f"Col{col_idx+1}",
        "dominant_type": dominant_type,
        "type_distribution": dict(type_counts.most_common(4)),
        "total_rows": total,
        "nulls": nulls,
        "null_pct": round(100 * nulls / total, 1) if total else 0,
        "unique_count": unique_count,
        "is_likely_categorical": unique_count < max(10, total * 0.05) and total > 20,
        "is_likely_unique_id": unique_count == len(non_empty) and len(non_empty) > 5,
        "samples": samples,
        "stats": stats if stats else None,
    }


def profile_worksheet(ws, max_sample_rows: int = 5) -> dict:
    """Profile a single openpyxl worksheet."""
    max_row = ws.max_row or 0
    max_col = ws.max_column or 0

    if max_row == 0 or max_col == 0:
        return {
            "sheet_name": ws.title,
            "dimensions": {"rows": 0, "cols": 0},
            "empty": True,
        }

    # Extract headers (row 1) and full column data (up to a reasonable limit for stats)
    # For very large sheets, sample rows for column profiling
    PROFILE_ROW_CAP = 5000  # Profile on first 5000 rows for speed

    headers = []
    for col_idx in range(1, max_col + 1):
        cell = ws.cell(row=1, column=col_idx)
        headers.append(str(cell.value).strip() if cell.value is not None else f"Col{col_idx}")

    # Extract column data (first PROFILE_ROW_CAP rows)
    profile_rows = min(max_row - 1, PROFILE_ROW_CAP)
    columns_data = [[] for _ in range(max_col)]
    for row_num in range(2, 2 + profile_rows):
        for col_num in range(1, max_col + 1):
            cell = ws.cell(row=row_num, column=col_num)
            columns_data[col_num - 1].append(cell.value)

    # Column profiles
    column_profiles = []
    for i, col_values in enumerate(columns_data):
        column_profiles.append(_column_profile(col_values, i, headers[i] if i < len(headers) else None))

    # Sample rows (first N + last N if sheet is large)
    first_rows = []
    for row_num in range(2, min(2 + max_sample_rows, max_row + 1)):
        row_vals = []
        for col_num in range(1, max_col + 1):
            v = ws.cell(row=row_num, column=col_num).value
            row_vals.append(str(v)[:40] if v is not None else "")
        first_rows.append(row_vals)

    last_rows = []
    if max_row > max_sample_rows * 2:
        for row_num in range(max(max_row - max_sample_rows + 1, max_sample_rows + 2), max_row + 1):
            row_vals = []
            for col_num in range(1, max_col + 1):
                v = ws.cell(row=row_num, column=col_num).value
                row_vals.append(str(v)[:40] if v is not None else "")
            last_rows.append(row_vals)

    # Formula count
    formula_count = 0
    for row in ws.iter_rows(min_row=1, max_row=min(max_row, 500), values_only=False):
        for cell in row:
            if cell.value is not None and isinstance(cell.value, str) and cell.value.startswith("="):
                formula_count += 1

    # Merged cells
    merged_ranges = [str(r) for r in ws.merged_cells.ranges][:10]

    return {
        "sheet_name": ws.title,
        "dimensions": {"rows": max_row, "cols": max_col, "profiled_rows": profile_rows + 1},
        "headers": headers,
        "columns": column_profiles,
        "sample_first_rows": first_rows,
        "sample_last_rows": last_rows,
        "formula_count_in_first_500_rows": formula_count,
        "merged_cell_ranges": merged_ranges,
        "truncated": max_row > PROFILE_ROW_CAP + 1,
    }


def profile_excel_file(file_path_or_bytes, max_sheets: int = 10) -> dict:
    """Profile an entire Excel workbook."""
    from openpyxl import load_workbook

    if isinstance(file_path_or_bytes, bytes):
        wb = load_workbook(io.BytesIO(file_path_or_bytes), data_only=False)
    else:
        wb = load_workbook(file_path_or_bytes, data_only=False)

    profile = {
        "file_type": "excel",
        "sheets": [],
        "sheet_count": len(wb.sheetnames),
        "defined_names": [],
    }

    # Defined names
    try:
        for name in wb.defined_names:
            try:
                dn = wb.defined_names[name]
                profile["defined_names"].append({
                    "name": name,
                    "value": str(dn.value)[:100] if hasattr(dn, 'value') else "",
                })
            except Exception:
                pass
    except Exception:
        pass

    for idx, sheet_name in enumerate(wb.sheetnames[:max_sheets]):
        try:
            ws = wb[sheet_name]
            profile["sheets"].append(profile_worksheet(ws))
        except Exception as e:
            profile["sheets"].append({"sheet_name": sheet_name, "error": str(e)})

    if len(wb.sheetnames) > max_sheets:
        profile["note"] = f"{len(wb.sheetnames) - max_sheets} additional sheets not profiled"

    return profile


def profile_to_llm_text(profile: dict, max_chars: int = 4000) -> str:
    """Convert a profile dict to condensed LLM-ready text.

    Target: 300-1500 tokens depending on sheet complexity.
    """
    parts = []
    parts.append(f"=== FILE PROFILE ({profile.get('file_type', 'unknown')}) ===")
    parts.append(f"Sheets: {profile.get('sheet_count', 0)}")

    if profile.get("defined_names"):
        dn_names = [d["name"] for d in profile["defined_names"][:5]]
        parts.append(f"Defined names: {', '.join(dn_names)}")

    for sheet in profile.get("sheets", []):
        if sheet.get("empty"):
            parts.append(f"\n--- Sheet: {sheet['sheet_name']} (EMPTY) ---")
            continue

        dim = sheet.get("dimensions", {})
        parts.append(f"\n--- Sheet: {sheet['sheet_name']} ({dim.get('rows', 0)} rows × {dim.get('cols', 0)} cols) ---")

        if sheet.get("merged_cell_ranges"):
            parts.append(f"Merged: {', '.join(sheet['merged_cell_ranges'][:5])}")
        if sheet.get("formula_count_in_first_500_rows", 0) > 0:
            parts.append(f"Formulas detected: {sheet['formula_count_in_first_500_rows']}+")

        # Column summaries — this is the most valuable part
        parts.append("Columns:")
        for col in sheet.get("columns", []):
            col_line = f"  [{chr(65 + col['column']) if col['column'] < 26 else 'C' + str(col['column']+1)}] {col['header'][:30]}"
            col_line += f" | type={col['dominant_type']}"
            col_line += f" | unique={col['unique_count']}"
            if col.get("null_pct", 0) > 0:
                col_line += f" | null={col['null_pct']}%"
            if col.get("is_likely_categorical"):
                col_line += " [CATEGORICAL]"
            if col.get("is_likely_unique_id"):
                col_line += " [UNIQUE_ID]"
            if col.get("stats"):
                s = col["stats"]
                col_line += f" | min={s.get('min', 0):.0f}, max={s.get('max', 0):.0f}, sum={s.get('sum', 0):.0f}"
            if col.get("samples"):
                samples_str = " | ".join(col["samples"][:2])
                col_line += f" | ex: {samples_str[:80]}"
            parts.append(col_line)

        # Sample rows — first 3 to orient the LLM
        if sheet.get("sample_first_rows"):
            parts.append("First rows:")
            for i, row in enumerate(sheet["sample_first_rows"][:3]):
                row_str = " | ".join(str(c)[:25] for c in row[:10])
                parts.append(f"  R{i+2}: {row_str}")

        if sheet.get("truncated"):
            parts.append(f"(profiled first {dim.get('profiled_rows', 0)} rows; sheet has {dim.get('rows', 0)} total rows)")

    text = "\n".join(parts)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n...[profile truncated]"
    return text


# === PDF / CSV / DOCX PROFILES ===

def profile_csv(file_path_or_bytes) -> dict:
    """Quick profile of a CSV file."""
    if isinstance(file_path_or_bytes, bytes):
        import csv, io
        text = file_path_or_bytes.decode('utf-8', errors='replace')
        reader = csv.reader(io.StringIO(text))
    else:
        import csv
        reader = csv.reader(open(file_path_or_bytes, 'r', encoding='utf-8', errors='replace'))

    rows = list(reader)
    if not rows:
        return {"file_type": "csv", "empty": True}

    headers = rows[0]
    data = rows[1:]
    max_col = len(headers)

    columns_data = [[] for _ in range(max_col)]
    for row in data[:5000]:
        for i in range(max_col):
            columns_data[i].append(row[i] if i < len(row) else "")

    column_profiles = [_column_profile(cd, i, headers[i] if i < len(headers) else None) for i, cd in enumerate(columns_data)]

    return {
        "file_type": "csv",
        "sheets": [{
            "sheet_name": "CSV",
            "dimensions": {"rows": len(rows), "cols": max_col, "profiled_rows": min(5000, len(data))},
            "headers": headers,
            "columns": column_profiles,
            "sample_first_rows": data[:5],
            "truncated": len(data) > 5000,
        }],
        "sheet_count": 1,
    }


def profile_pdf(file_path_or_bytes, max_pages: int = 5) -> dict:
    """Quick profile of a PDF — first few pages + tables."""
    try:
        import pdfplumber
        if isinstance(file_path_or_bytes, bytes):
            pdf = pdfplumber.open(io.BytesIO(file_path_or_bytes))
        else:
            pdf = pdfplumber.open(file_path_or_bytes)

        total_pages = len(pdf.pages)
        text_sample = ""
        tables_detected = 0
        for page in pdf.pages[:max_pages]:
            text_sample += (page.extract_text() or "")[:2000] + "\n\n"
            try:
                tables = page.extract_tables()
                tables_detected += len(tables)
            except Exception:
                pass

        return {
            "file_type": "pdf",
            "total_pages": total_pages,
            "profiled_pages": min(max_pages, total_pages),
            "tables_detected": tables_detected,
            "text_sample": text_sample[:3000],
        }
    except Exception as e:
        return {"file_type": "pdf", "error": str(e)}


def profile_docx(file_path_or_bytes) -> dict:
    """Quick profile of a DOCX."""
    try:
        import docx
        if isinstance(file_path_or_bytes, bytes):
            doc = docx.Document(io.BytesIO(file_path_or_bytes))
        else:
            doc = docx.Document(file_path_or_bytes)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = doc.tables

        table_summaries = []
        for t in tables[:3]:
            rows = []
            for row in t.rows[:5]:
                cells = [c.text[:30] for c in row.cells]
                rows.append(cells)
            table_summaries.append({"rows": len(t.rows), "cols": len(t.columns), "sample": rows})

        return {
            "file_type": "docx",
            "paragraphs": len(paragraphs),
            "tables": len(tables),
            "first_paragraphs": paragraphs[:10],
            "table_summaries": table_summaries,
        }
    except Exception as e:
        return {"file_type": "docx", "error": str(e)}


# === MAIN ENTRY POINT ===

def profile_file(filename: str, content: bytes) -> dict:
    """Auto-detect file type and return appropriate profile."""
    fn = filename.lower()
    try:
        if fn.endswith((".xlsx", ".xlsm")):
            return profile_excel_file(content)
        elif fn.endswith(".csv"):
            return profile_csv(content)
        elif fn.endswith(".pdf"):
            return profile_pdf(content)
        elif fn.endswith(".docx"):
            return profile_docx(content)
        else:
            return {"file_type": "unknown", "size": len(content)}
    except Exception as e:
        logger.warning(f"Profile failed for {filename}: {e}")
        return {"file_type": "unknown", "error": str(e), "size": len(content)}


def profile_uploaded_files(uploaded_files: list[dict], max_total_chars: int = 8000) -> str:
    """Profile all uploaded files and return a combined LLM-ready text."""
    profiles = []
    for f in uploaded_files:
        filename = f.get("filename", "unknown")
        content = f.get("content", b"")
        profile = profile_file(filename, content)
        profiles.append((filename, profile))

    parts = [f"=== UPLOADED FILES PROFILES ({len(profiles)} file(s)) ===\n"]
    per_file_budget = max_total_chars // max(len(profiles), 1)

    for filename, profile in profiles:
        parts.append(f"\n📁 FILE: {filename}")
        if profile.get("file_type") in ("excel", "csv"):
            parts.append(profile_to_llm_text(profile, max_chars=per_file_budget))
        elif profile.get("file_type") == "pdf":
            parts.append(f"  PDF: {profile.get('total_pages', 0)} pages, {profile.get('tables_detected', 0)} tables detected")
            parts.append(f"  Text sample: {(profile.get('text_sample', '')[:per_file_budget-200])}")
        elif profile.get("file_type") == "docx":
            parts.append(f"  DOCX: {profile.get('paragraphs', 0)} paragraphs, {profile.get('tables', 0)} tables")
            parts.append(f"  First paragraphs: {' | '.join(profile.get('first_paragraphs', [])[:5])[:per_file_budget-200]}")
        else:
            parts.append(f"  [Unsupported or profile error: {profile.get('error', 'unknown')}]")

    return "\n".join(parts)

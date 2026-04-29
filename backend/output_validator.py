"""
Output File Validator — sanity checks on files produced by the code executor.

For every file the agent produces, we verify:
- Excel files: open without error, have non-empty sheets, formulas are preserved
- DOCX files: valid structure, non-empty body
- PDF files: not corrupted, has at least one page
- Images: valid format, reasonable dimensions

Catches cases where code runs successfully but produces broken output
(e.g., empty Excel, corrupted DOCX, PDF with 0 pages).
"""

import io
import logging
from typing import Optional

logger = logging.getLogger("output_validator")


def validate_excel(content: bytes) -> dict:
    """Validate an Excel file. Returns {valid, sheets, rows, formulas_count, issues}."""
    result = {"valid": False, "sheets": 0, "rows": 0, "formulas_count": 0, "issues": []}
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), data_only=False)
        result["sheets"] = len(wb.sheetnames)
        if result["sheets"] == 0:
            result["issues"].append("No sheets in workbook")
            return result

        total_rows = 0
        formulas = 0
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            total_rows += ws.max_row
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None and isinstance(cell.value, str) and cell.value.startswith("="):
                        formulas += 1
        result["rows"] = total_rows
        result["formulas_count"] = formulas

        if total_rows == 0:
            result["issues"].append("All sheets empty")
            return result

        result["valid"] = True
    except Exception as e:
        result["issues"].append(f"Excel validation error: {type(e).__name__}: {e}")
    return result


def validate_docx(content: bytes) -> dict:
    """Validate a DOCX file. Returns {valid, paragraphs, tables, total_chars, issues}."""
    result = {"valid": False, "paragraphs": 0, "tables": 0, "total_chars": 0, "issues": []}
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    total_chars += len(cell.text)

        result["paragraphs"] = para_count
        result["tables"] = table_count
        result["total_chars"] = total_chars

        if total_chars < 50:
            result["issues"].append(f"Document too short ({total_chars} chars) — likely empty or skeleton")
            return result

        result["valid"] = True
    except Exception as e:
        result["issues"].append(f"DOCX validation error: {type(e).__name__}: {e}")
    return result


def validate_pdf(content: bytes) -> dict:
    """Validate a PDF file. Returns {valid, pages, issues}."""
    result = {"valid": False, "pages": 0, "issues": []}
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        result["pages"] = len(reader.pages)
        if result["pages"] == 0:
            result["issues"].append("PDF has 0 pages")
            return result
        result["valid"] = True
    except Exception as e:
        result["issues"].append(f"PDF validation error: {type(e).__name__}: {e}")
    return result


def validate_image(content: bytes) -> dict:
    """Validate an image (JPG/PNG). Returns {valid, format, size, issues}."""
    result = {"valid": False, "format": "", "size": (0, 0), "issues": []}
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(content))
        result["format"] = img.format or ""
        result["size"] = img.size
        if img.size[0] < 10 or img.size[1] < 10:
            result["issues"].append(f"Image too small: {img.size}")
            return result
        result["valid"] = True
    except Exception as e:
        result["issues"].append(f"Image validation error: {type(e).__name__}: {e}")
    return result


def validate_file(filename: str, content: bytes) -> dict:
    """Validate any output file based on extension."""
    fn_lower = filename.lower()
    if len(content) == 0:
        return {"valid": False, "type": "empty", "issues": ["File is empty (0 bytes)"]}

    if fn_lower.endswith((".xlsx", ".xlsm")):
        r = validate_excel(content)
        r["type"] = "excel"
        return r
    elif fn_lower.endswith(".docx"):
        r = validate_docx(content)
        r["type"] = "docx"
        return r
    elif fn_lower.endswith(".pdf"):
        r = validate_pdf(content)
        r["type"] = "pdf"
        return r
    elif fn_lower.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp")):
        r = validate_image(content)
        r["type"] = "image"
        return r
    elif fn_lower.endswith((".csv", ".txt", ".html", ".json", ".xml")):
        # Simple text check
        try:
            text = content.decode("utf-8", errors="replace")
            return {"valid": len(text.strip()) > 0, "type": "text",
                    "chars": len(text), "issues": [] if len(text.strip()) > 0 else ["Empty text file"]}
        except Exception as e:
            return {"valid": False, "type": "text", "issues": [str(e)]}
    else:
        # Unknown type — just check non-empty
        return {"valid": len(content) > 0, "type": "unknown",
                "size": len(content), "issues": [] if len(content) > 0 else ["Empty"]}


def validate_all_outputs(output_files: list[dict]) -> dict:
    """Validate a list of output files from the code executor.

    Input: [{"name": ..., "content_b64": ..., "size": ...}, ...]
    Returns: {"all_valid": bool, "results": [{name, valid, type, ...}], "summary": str}
    """
    import base64
    results = []
    all_valid = True
    for f in output_files:
        try:
            content = base64.b64decode(f.get("content_b64", ""))
        except Exception as e:
            results.append({"name": f.get("name", "unknown"), "valid": False, "issues": [f"Base64 decode failed: {e}"]})
            all_valid = False
            continue

        v = validate_file(f.get("name", "unknown"), content)
        v["name"] = f.get("name", "unknown")
        results.append(v)
        if not v.get("valid"):
            all_valid = False

    # Summary
    valid_count = sum(1 for r in results if r.get("valid"))
    summary_parts = [f"{valid_count}/{len(results)} files valid"]
    for r in results:
        if not r.get("valid"):
            summary_parts.append(f"INVALID: {r['name']} — {'; '.join(r.get('issues', []))}")
    summary = "\n".join(summary_parts)

    return {"all_valid": all_valid, "results": results, "summary": summary}

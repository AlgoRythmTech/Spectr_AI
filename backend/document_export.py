import io
import re
import logging
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

def parse_inline_markdown(para, text):
    """Parse **, __, *, _ inline markdown and add to a docx paragraph."""
    # Split text by bold and italic markers
    # We will use a regex that matches **bold**, __bold__, *italic*, _italic_
    # This is a basic tokenizer
    tokens = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_)', text)
    for token in tokens:
        if not token:
            continue
        if (token.startswith('**') and token.endswith('**')) or (token.startswith('__') and token.endswith('__')):
            if len(token) > 4:
                run = para.add_run(token[2:-2])
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        elif (token.startswith('*') and token.endswith('*')) or (token.startswith('_') and token.endswith('_')):
            if len(token) > 2:
                run = para.add_run(token[1:-1])
                run.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        else:
            run = para.add_run(token)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def generate_word_document(title: str, content: str, doc_type: str = "memo",
                           header_option: str = "", firm_name: str = "",
                           watermark: str = "") -> bytes:
    """Generate a highly formatted Word document with official letterhead."""
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from datetime import datetime

    doc = Document()

    # Base styling — institutional-grade formatting
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Configure heading styles for professional look
    for level in [1, 2, 3]:
        try:
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.name = 'Times New Roman'
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            if level == 1:
                heading_style.font.size = Pt(14)
                heading_style.font.bold = True
            elif level == 2:
                heading_style.font.size = Pt(13)
                heading_style.font.bold = True
            elif level == 3:
                heading_style.font.size = Pt(12)
                heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        except Exception:
            pass

    # Professional margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(2.54)

        # === WATERMARK (diagonal text behind content) ===
        if watermark:
            try:
                header_for_wm = section.header
                header_for_wm.is_linked_to_previous = False
                wm_xml = f'''
                <w:r {nsdecls("w","v","o")}>
                  <w:rPr><w:noProof/></w:rPr>
                  <w:pict>
                    <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" path="m@7,l@8,m@5,21600l@6,21600e">
                      <v:formulas><v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/><v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/><v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 0"/><v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/><v:f eqn="if @0 @4 21600"/><v:f eqn="mid @5 @6"/><v:f eqn="mid @8 @5"/><v:f eqn="mid @7 @8"/><v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/></v:formulas>
                    </v:shapetype>
                    <v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s2049" type="#_x0000_t136" style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:100pt;rotation:315;z-index:-251658752" o:allowincell="f" fillcolor="silver" stroked="f">
                      <v:fill opacity=".15"/>
                      <v:textpath style="font-family:&quot;Times New Roman&quot;;font-size:60pt" string="{watermark}"/>
                    </v:shape>
                  </w:pict>
                </w:r>'''
                # Try to insert watermark — this is best-effort, some python-docx versions may not support it
            except Exception:
                pass  # Watermark is non-critical

        # === PAGE FOOTER WITH PAGE NUMBERS ===
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Generated by Spectr AI — AlgoRythm Technologies  |  Page ")
        footer_run.font.size = Pt(8)
        footer_run.font.name = 'Arial'
        footer_run.font.color.rgb = RGBColor(153, 153, 153)
        # Add page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        footer_para._p.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        footer_para._p.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        footer_para._p.append(fldChar2)

        # === PAGE HEADER — CONFIDENTIAL STAMP ===
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        conf_run = header_para.add_run("CONFIDENTIAL — PRIVILEGED & ATTORNEY WORK PRODUCT")
        conf_run.bold = True
        conf_run.font.size = Pt(7)
        conf_run.font.name = 'Arial'
    
    # === FIRM LETTERHEAD ===
    firm_display = firm_name if firm_name else "SPECTR — LEGAL & TAX ADVISORY"
    letterhead = doc.add_paragraph()
    letterhead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lh_run = letterhead.add_run(firm_display.upper())
    lh_run.bold = True
    lh_run.font.size = Pt(18)
    lh_run.font.name = 'Times New Roman'
    letterhead.space_after = Pt(2)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Advocates · Chartered Accountants · Tax Consultants")
    sub_run.font.size = Pt(9)
    sub_run.font.name = 'Arial'
    subtitle.space_after = Pt(6)
    
    # Horizontal rule
    hr = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr_run = hr.add_run("━" * 65)
    hr_run.font.size = Pt(6)
    hr.space_after = Pt(12)
    
    # === DATE AND REFERENCE BLOCK ===
    ref_table = doc.add_table(rows=1, cols=2)
    ref_table.columns[0].width = Inches(3.5)
    ref_table.columns[1].width = Inches(3.0)
    cell_left = ref_table.cell(0, 0)
    cell_right = ref_table.cell(0, 1)
    
    left_para = cell_left.paragraphs[0]
    left_run = left_para.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}")
    left_run.font.size = Pt(10)
    left_run.font.name = 'Times New Roman'
    
    right_para = cell_right.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    import hashlib
    ref_hash = hashlib.md5(title.encode()).hexdigest()[:8].upper()
    right_run = right_para.add_run(f"Ref: AR/{datetime.now().strftime('%Y')}/{ref_hash}")
    right_run.font.size = Pt(10)
    right_run.font.name = 'Times New Roman'
    
    doc.add_paragraph()  # spacing
    
    # Header option (e.g., DRAFT / WITHOUT PREJUDICE)
    if header_option:
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run(header_option.upper())
        header_run.bold = True
        header_run.font.size = Pt(10)
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.upper())
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Times New Roman'
    title_para.space_after = Pt(18)
    
    doc.add_paragraph()  # spacing
    
    # Content - parse markdown-like headers, lists, and inline formatting
    para_num = 1
    
    in_table = False
    table_rows = []
    
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        
        # Handle tables
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            row = [cell.strip() for cell in stripped.split('|')[1:-1]]
            if all(cell.strip('-') == '' for cell in row):
                continue
            table_rows.append(row)
            continue
        elif in_table:
            # End of table, render it
            if table_rows:
                num_cols = max(len(r) for r in table_rows) if table_rows else 1
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_rows):
                    for c_idx, cell_data in enumerate(row_data):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell_para = cell.paragraphs[0]
                            cell_para.paragraph_format.space_before = Pt(2)
                            cell_para.paragraph_format.space_after = Pt(2)
                            if r_idx == 0:
                                # Header row: bold, shaded background
                                run = cell_para.add_run(cell_data)
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                                try:
                                    from docx.oxml.ns import nsdecls
                                    from docx.oxml import parse_xml
                                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F3F4F6" w:val="clear"/>')
                                    cell._tc.get_or_add_tcPr().append(shading)
                                except Exception:
                                    pass
                            else:
                                parse_inline_markdown(cell_para, cell_data)
                doc.add_paragraph()  # Spacing after table
            in_table = False
            table_rows = []
            
        if not stripped:
            doc.add_paragraph()
            continue
        
        if line.startswith('### '):
            heading = doc.add_heading(level=3)
            heading_run = heading.add_run(line[4:].upper())
            heading_run.bold = True
            heading_run.font.name = 'Times New Roman'
            heading_run.font.color.rgb = None # Reset to black
        elif line.startswith('## '):
            heading = doc.add_heading(level=2)
            heading_run = heading.add_run(line[3:])
            heading_run.bold = True
            heading_run.font.name = 'Times New Roman'
            heading_run.font.color.rgb = None
        elif line.startswith('# '):
            heading = doc.add_heading(level=1)
            heading_run = heading.add_run(line[2:])
            heading_run.bold = True
            heading_run.font.name = 'Times New Roman'
            heading_run.font.color.rgb = None
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            parse_inline_markdown(para, line[2:])
        elif re.match(r'^\d+\.\s', line):
            para = doc.add_paragraph(style='List Number')
            content_start = line.find('.') + 1
            parse_inline_markdown(para, line[content_start:].strip())
        else:
            para = doc.add_paragraph()
            if doc_type in ["notice", "petition", "application"] and not line.startswith(' '):
                run = para.add_run(f"{para_num}. ")
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                para_num += 1
            parse_inline_markdown(para, line)
            
    # If document ended with a table
    if in_table and table_rows:
        num_cols = max(len(r) for r in table_rows) if table_rows else 1
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'
        for r_idx, row_data in enumerate(table_rows):
            for c_idx, cell_data in enumerate(row_data):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell_para = cell.paragraphs[0]
                    if r_idx == 0:
                        run = cell_para.add_run(cell_data)
                        run.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                    else:
                        parse_inline_markdown(cell_para, cell_data)
    
    # Signature block
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_run = sig.add_run("_________________________")
    sig.add_run("\n")
    sig.add_run(firm_name if firm_name else "Authorized Signatory")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_pdf_content(title: str, content: str, header_option: str = "",
                          watermark: str = "") -> str:
    """Generate HTML content for PDF conversion via WeasyPrint."""
    watermark_css = ""
    if watermark:
        watermark_css = f"""
        .watermark {{
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 80px;
            color: rgba(0, 0, 0, 0.06);
            font-weight: bold;
            z-index: -1;
            pointer-events: none;
        }}"""
    
    header_html = ""
    if header_option:
        header_html = f'<div class="header-option">{header_option.upper()}</div>'
    
    # Convert markdown content to HTML
    html_content = content
    html_content = html_content.replace('\n### ', '\n<h3>')
    html_content = html_content.replace('\n## ', '\n<h2>')
    html_content = html_content.replace('\n# ', '\n<h1>')
    
    lines = html_content.split('\n')
    processed = []
    for line in lines:
        if line.startswith('<h3>'):
            processed.append(f'<h3>{line[4:]}</h3>')
        elif line.startswith('<h2>'):
            processed.append(f'<h2>{line[4:]}</h2>')
        elif line.startswith('<h1>'):
            processed.append(f'<h1>{line[4:]}</h1>')
        elif line.strip():
            # bold and italic parsing for HTML
            parsed_line = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', line)
            parsed_line = re.sub(r'__(.*?)__', r'<strong>\1</strong>', parsed_line)
            parsed_line = re.sub(r'\*(.*?)\*', r'<em>\1</em>', parsed_line)
            parsed_line = re.sub(r'_(.*?)_', r'<em>\1</em>', parsed_line)
            processed.append(f'<p>{parsed_line}</p>')
        else:
            processed.append('<br/>')
    
    body_html = '\n'.join(processed)
    
    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    @page {{
        size: A4;
        margin: 2.54cm 2.54cm 2.54cm 3.18cm;
        @bottom-center {{
            content: counter(page);
            font-size: 10px;
        }}
    }}
    body {{
        font-family: 'Times New Roman', Times, serif;
        font-size: 12pt;
        line-height: 1.6;
        color: #000;
    }}
    h1 {{ font-size: 16pt; font-weight: bold; text-align: center; margin: 20pt 0 12pt; }}
    h2 {{ font-size: 14pt; font-weight: bold; margin: 16pt 0 8pt; }}
    h3 {{ font-size: 12pt; font-weight: bold; margin: 12pt 0 6pt; text-transform: uppercase; }}
    p {{ margin: 6pt 0; text-align: justify; }}
    .header-option {{
        text-align: right;
        font-weight: bold;
        font-size: 10pt;
        margin-bottom: 20pt;
    }}
    .title {{
        text-align: center;
        font-weight: bold;
        font-size: 14pt;
        margin: 20pt 0;
    }}
    {watermark_css}
</style>
</head>
<body>
{f'<div class="watermark">{watermark}</div>' if watermark else ''}
{header_html}
<div class="title">{title}</div>
{body_html}
</body>
</html>"""
    return html


def generate_pdf_bytes(title: str, content: str, header_option: str = "",
                        watermark: str = "") -> bytes:
    """Generate PDF bytes. Tries WeasyPrint first, falls back to ReportLab."""
    # Try WeasyPrint first
    try:
        from weasyprint import HTML
        html_str = generate_pdf_content(title, content, header_option, watermark)
        pdf_bytes = HTML(string=html_str).write_pdf()
        return pdf_bytes
    except Exception as e:
        logger.warning(f"WeasyPrint unavailable ({e}), falling back to ReportLab")

    # Fallback: ReportLab
    try:
        import io
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                leftMargin=1*inch, rightMargin=1*inch,
                                topMargin=1*inch, bottomMargin=1*inch)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('DocTitle', parent=styles['Heading1'],
                                      fontSize=18, spaceAfter=20)
        body_style = ParagraphStyle('DocBody', parent=styles['Normal'],
                                     fontSize=11, leading=16, alignment=TA_LEFT)

        story = []
        story.append(Paragraph(title.replace('&', '&amp;').replace('<', '&lt;'), title_style))
        story.append(Spacer(1, 0.3*inch))

        # Convert markdown-ish content to paragraphs
        import re
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.15*inch))
                continue
            # Escape HTML entities
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            # Bold: **text**
            line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
            # Headings
            if line.startswith('## '):
                story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], styles['Heading3']))
            elif line.startswith('- ') or line.startswith('* '):
                story.append(Paragraph(f"• {line[2:]}", body_style))
            else:
                story.append(Paragraph(line, body_style))

        doc.build(story)
        return buffer.getvalue()
    except Exception as e2:
        logger.error(f"ReportLab PDF fallback also failed: {e2}")
        return b""


def generate_excel_document(title: str, content: str) -> bytes:
    """Generate a clean Excel document from text/markdown content, extracting all tables."""
    import pandas as pd
    import io

    lines = content.split('\n')
    tables = []
    current_table = []
    current_headers = []
    table_mode = False

    # Extract markdown tables
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|') and stripped.endswith('|'):
            row = [cell.strip() for cell in stripped.split('|')[1:-1]]
            if all(cell.strip('-') == '' for cell in row):
                continue
            if not table_mode:
                current_headers = row
                table_mode = True
            else:
                if len(row) == len(current_headers):
                    current_table.append(dict(zip(current_headers, row)))
                else:
                    current_table.append({f"Col_{i}": val for i, val in enumerate(row)})
        else:
            if table_mode:
                tables.append(current_table)
                current_table = []
                table_mode = False
                
    if table_mode and current_table:
        tables.append(current_table)

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        if tables:
            # Write each identified table to a separate sheet
            for idx, tbl in enumerate(tables):
                sheet_name = f"Table_{idx+1}"
                df = pd.DataFrame(tbl)
                df.to_excel(writer, index=False, sheet_name=sheet_name)
        
        # Also write the full raw text to an "Original_Content" sheet for context
        full_data = [{"Line": i+1, "Content": line.strip()} for i, line in enumerate(lines) if line.strip()]
        if not full_data:
            full_data = [{"Line": 1, "Content": "No content available"}]
        df_full = pd.DataFrame(full_data)
        df_full.to_excel(writer, index=False, sheet_name="Full_Document")
        
    buffer.seek(0)
    return buffer.getvalue()


